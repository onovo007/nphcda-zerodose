"""Generate (1) the dropout-driver direction+CI chart (validates the LASSO importance slide) and
(2) a one-page presenter run-sheet (.docx). Saves into 02_PPT_Slide-Deck/."""
import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

OUT = ("C:/Users/Amobi Andrew/OneDrive/Desktop/Immunization Modeling Fellowship/In-person Training/"
       "Submission Package/ReadyOne/All I Need/02_PPT_Slide-Deck/")
GREEN, NAVY, RED = "#1C7A3D", "#1F3B57", "#C0392B"


def build_driver_direction():
    import data_io as io
    from models.d2_dropout import lasso_drivers_inference
    s = io.load_sample()
    agg = io.state_monthly(io.prep_dhis2(s["dhis2"]))
    res = lasso_drivers_inference(agg, s["model_dataset"], key="png-dir", n_boot=120)
    fig, axes = plt.subplots(1, 3, figsize=(13.333, 6), dpi=200)
    for ax, (pair, df) in zip(axes, res.items()):
        df = df.iloc[::-1].reset_index(drop=True)  # top driver at top
        coef = df["Std coef"].astype(float).values
        lo, hi = [], []
        for ci in df["95% CI"]:
            a, b = ci.split(" to "); lo.append(float(a)); hi.append(float(b))
        lo, hi = np.array(lo), np.array(hi)
        y = np.arange(len(df))
        for i in range(len(df)):
            c = RED if coef[i] > 0 else GREEN
            ax.plot([lo[i], hi[i]], [y[i], y[i]], color=c, lw=2.4, solid_capstyle="round")
            ax.plot(coef[i], y[i], "o", color=c, ms=8, zorder=3)
        ax.axvline(0, color="#888", ls="--", lw=1.2)
        ax.set_yticks(y); ax.set_yticklabels(df["Driver"], fontsize=8.5)
        ax.set_title(pair, fontsize=11.5, fontweight="bold", color=NAVY)
        ax.set_xlabel("Std. coefficient (95% CI)", fontsize=9)
        for sp in ("top", "right"):
            ax.spines[sp].set_visible(False)
    fig.suptitle("Drivers of dropout: direction and uncertainty, not just importance",
                 fontsize=17, fontweight="bold", color=NAVY, y=1.02)
    fig.text(0.5, -0.02, "Right of 0 = increases dropout (red); left of 0 = reduces it (green); bars "
             "crossing 0 are uncertain. Top-4 LASSO-selected drivers per transition; ecological "
             "associations across all 37 states (a census), directional not causal.",
             ha="center", fontsize=9, color="#5B6B79")
    fig.tight_layout()
    fig.savefig(OUT + "NPHCDA_Dropout_Drivers_Direction_CI.png", bbox_inches="tight", facecolor="white")
    plt.close(fig); print("WROTE driver-direction chart")


def build_runsheet():
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    navy, green, mute = RGBColor(0x1F, 0x3B, 0x57), RGBColor(0x1C, 0x7A, 0x3D), RGBColor(0x6B, 0x7A, 0x88)
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.5)
        sec.left_margin = sec.right_margin = Inches(0.6)
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(9.5)

    def head(t, sz=13, col=navy):
        p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.size = Pt(sz); r.font.color.rgb = col
        p.space_after = Pt(3); return p

    def tbl(headers, rows, widths=None):
        t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
        for i, h in enumerate(headers):
            rr = t.rows[0].cells[i].paragraphs[0].add_run(h); rr.bold = True; rr.font.size = Pt(9)
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = str(v)
                for pp in cells[i].paragraphs:
                    for rn in pp.runs:
                        rn.font.size = Pt(8.5)
        return t

    h = doc.add_paragraph(); r = h.add_run("Presenter run-sheet - NPHCDA Zero-Dose review")
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = navy
    s = doc.add_paragraph(); sr = s.add_run("20 minutes slides + 10 minutes live web app")
    sr.italic = True; sr.font.size = Pt(10); sr.font.color.rgb = green

    head("Slides - 20 minutes")
    tbl(["Time", "Slides", "Land this"], [
        ["0:00-0:30", "Title", "NPHCDA-owned, GAVI/UNICEF-supported, delivered as a live tool"],
        ["0:30-2:30", "Exec summary", "Flat rate hides an emergency (~2.07M); North-West concentrated; coverage gaps = death gaps; fix is local"],
        ["2:30-3:30", "The problem", "A decade flat, off the IA2030 path"],
        ["3:30-5:00", "Scope/Methods/Data", "Co-designed, reproducible; denominator anchored to UN WPP 2024 (~7.5M births); live births = sensitivity"],
        ["5:00-6:30", "Domain 1", "National OK but 1,645 LGA-antigen breaches across 591 LGAs - risk is local"],
        ["6:30-8:00", "Domain 2", "Three transitions, three different drivers -> three tailored responses"],
        ["8:00-9:00", "Domain 3/4", "Equity gradient; demand != supply (rho=-0.33)"],
        ["9:00-14:00", "Domain 5 (core)", "~2.09M in 2026; NW ~1.08M; top 20% of LGAs = 62% of burden / 80% in ~270 LGAs; 5 archetypes"],
        ["14:00-15:30", "Domain 6/7", "Independent data converge on same states; coverage tracks deaths (r=0.87)"],
        ["15:30-17:00", "Convergence", "Four independent analyses -> same place: strongest corroboration short of a trial"],
        ["17:00-19:00", "Means/Limits/Next", "4 design implications; honest limitations; next steps incl. DIH handover + platform"],
        ["19:00-20:00", "Bridge", "'Everything is live and reproducible - let me show you' -> open the app"],
    ])

    head("Web app - 10 minutes")
    tbl(["Time", "Page", "Do / say"], [
        ["0:00-0:40", "Home", "No-code, live on Hugging Face; click Use bundled sample data"],
        ["0:40-1:40", "Data & Quality", "Reporting heatmap -> By-LGA drill-down + Anomaly tab"],
        ["1:40-3:40", "Coverage Forecasting", "Toggle WHO admin coverage -> denominator -> NDHS survey lines; change 3/6/12m; LGA at-risk screen"],
        ["3:40-4:30", "Dropout & Completion", "Drivers with direction + 95% CI (uncertainty, not just importance)"],
        ["4:30-7:00", "Zero-Dose & Hotspots", "Bayesian forecast + CIs -> LGA Pareto -> Gi* map (bold boundaries) -> run leave-one-wave-out"],
        ["7:00-8:00", "Implementation Science", "Run a hypothesis test or multiple regression live"],
        ["8:00-9:00", "Ask the Analyst (ZARA)", "Ask: which states/LGAs to prioritize first, and why"],
        ["9:00-10:00", "Reports & Briefs", "Generate the GAVI factsheet; show Word/PowerPoint/Methods downloads"],
    ])

    head("Formulas (for Q&A)", 12)
    for line in [
        "Zero-dose rate = (12-23m children with no Penta1) / (all 12-23m children) x 100; modelled on the logit scale (Bayesian hierarchical Beta).",
        "Zero-dose burden (state) = zero-dose rate(2026) x infant cohort; infant cohort = under-five population / 5.",
        "LGA: cohort_LGA = cohort_state x (LGA pop / sum of state LGA pop); burden_LGA = LGA rate x cohort_LGA; Pareto ranks LGAs by burden.",
        "Coverage - index = forecast monthly doses / mean 2024 monthly doses x 100 (denominator-free; 80% line = at-risk-of-decline early-warning).",
        "Coverage - WHO admin = monthly doses / (annual eligible infants / 12) x 100; denominator = surviving infants (under-five/5 /12, or DHIS2 live births/12).",
        "Dropout(A->B) = (doses_A - doses_B) / doses_A x 100.",
    ]:
        p = doc.add_paragraph(line, style="List Bullet")
        for rn in p.runs:
            rn.font.size = Pt(8.8)

    head("Why under-five / 5 (denominator)", 12)
    p = doc.add_paragraph(
        "The under-five population covers five single-year age bands (0-4 years); dividing by five gives "
        "the average size of one single-year cohort - our proxy for the 12-23-month (surviving-infant) "
        "denominator. It is a recognized proxy when a direct single-year denominator is unavailable; the "
        "gold standard is surviving infants (live births minus infant deaths). Our proxy lands at ~7.0M, "
        "consistent with UN World Population Prospects 2024 (~7.51M Nigerian births in 2023, via Our World "
        "in Data; minus infant mortality ~= ~7.0M surviving infants) - so it is empirically anchored. We "
        "used the 2024 figure (more stable than 2025) and expose DHIS2 live births (~2.5M, facility-"
        "reported, under-counted) only as a sensitivity check; official NPC surviving-infant projections "
        "can be substituted directly in the app.")
    for rn in p.runs:
        rn.font.size = Pt(8.8)

    doc.save(OUT + "NPHCDA_Presenter_Run_Sheet.docx")
    print("WROTE run-sheet")


build_driver_direction()
build_runsheet()
print("DONE")
